import docx


def get_obj(app,name,fathername,who, who_name,surname):
    return {
        "{{app}}": app,
        "{{name}}": name,
        "{{fathername}}": fathername,
        "who": who,
        "who_name": who_name,
        "surname": surname
    }

res = []
with open("data.txt", "r", encoding="utf-8") as f:
    lines = f.read().split("\n")
    for line in lines:
        res.append(get_obj(*line.split(",")))



from docx.shared import Pt


for i in res:
    doc = docx.Document('1.docx')

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    for k,v in i.items():
        if "{{" in k:
            doc.paragraphs[6].text = doc.paragraphs[6].text.replace(k,v)

    doc.paragraphs[6].style = style

    doc.tables[1].rows[0].cells[1].text = f'{i["who"]}\n\n{i["who_name"]}'

    doc.save(f"{i['surname']} {i['{{name}}']} {i['{{fathername}}']}.docx")
